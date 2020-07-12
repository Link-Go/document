import os

from datetime import datetime
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from fastapi import FastAPI
from fastapi.responses import FileResponse

from utils import set_cell_border

app = FastAPI()

# overall_path = "./guide/全局分析指引.docx"
artificial_template_path = "./guide/人工核查指引.docx"
artificial_docx_path = "./guide/artificial"


async def download_path(kind: str, cover_data: List[str],
                        content_data: List[List[str]]) -> Tuple[str, str]:
    """生成docx文档并存在本地"""
    # 获取模板
    document = Document(artificial_template_path)
    # 指定字体格式，大小
    style = document.styles['Normal']
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u"微软雅黑")
    style.font.size = Pt(10)

    # 封面表格，内容表格
    table_cover, table_content = document.tables[0], document.tables[1]

    # 更新到封面表格
    for i, row in enumerate(table_cover.rows[:]):
        row.cells[1].text = cover_data[i]

    # 删除模板中示例数据
    for _, row in enumerate(table_content.rows[1:]):
        row._element.getparent().remove(row._element)

    # 更新到内容表格
    for i in range(len(content_data)):
        row_cells = table_content.add_row().cells
        for j, cell in enumerate(row_cells):
            cell.text = content_data[i][j]
            # TODO 增加是否为高风险项的判断
            if j == 1:
                # 添加图片
                run = cell.add_paragraph().add_run()
                run.add_picture("./guide/高风险项.png", width=Inches(0.6))
                # 添加文字
                # p = cell.add_paragraph()
                # run = p.add_run("高风险")
                # run.font.size = Pt(10)
                # run.font.color.rgb = RGBColor(255, 0, 0)
            # if j == 1:
            #     table = cell.add_table(1, 1, )
            #     shading_elm_1 = parse_xml(
            #         r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
            #     table.rows[0].cells[0]._tc.get_or_add_tcPr().append(
            #         shading_elm_1)
            #     table.cell(0, 0).width = Inches(1)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = Pt(18)

    if not os.path.exists(artificial_docx_path):
        os.mkdir(artificial_docx_path)

    # 文件名以 人工核查指引 + 时间 命名
    time_format = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = kind + time_format + r".docx"
    file_path = artificial_docx_path + "/" + file_name
    document.save(file_path)

    return (file_name, file_path)


@app.get("/")
async def download():
    """下载配置指引"""
    # TODO 从数据获取数据，并处理成如下格式
    cover_data = ["财务系统", "第三级", "2020-06-05 12:30:06"]
    content_data = [[
        "安全通信网络-网络架构", "应采用校验技术或密码技术保证通信过程中数据的完整 ", "防火墙",
        "确认业务系统是否装有防火墙：若已安装，检测是否配置相关策略，若未安装建议购买安装", "", ""
    ],
                    [
                        "安全通信网络-网络架构", "应采用校验技术或密码技术保证通信过程中数据的完整", "防火墙",
                        "确认业务系统是否装有防火墙：若已安装，检测是否配置相关策略，若未安装建议购买安装", "", ""
                    ],
                    [
                        "安全通信网络-网络架构", "应采用校验技术或密码技术保证通信过程中数据的完整", "防火墙",
                        "确认业务系统是否装有防火墙：若已安装，检测是否配置相关策略，若未安装建议购买安装", "", ""
                    ],
                    [
                        "安全通信网络-网络架构", "应采用校验技术或密码技术保证通信过程中数据的完整", "防火墙",
                        "确认业务系统是否装有防火墙：若已安装，检测是否配置相关策略，若未安装建议购买安装", "", ""
                    ]]

    file_name, file_path = await download_path("人工核查指引", cover_data,
                                               content_data)
    return FileResponse(file_path, filename=file_name)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("demo:app", reload=True)
